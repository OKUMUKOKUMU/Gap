import streamlit as st
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import io
from datetime import datetime

def create_word_report(data):
    doc = Document()
    
    # Title
    title = doc.add_paragraph()
    title_run = title.add_run(f"Weekly Sales Report - Week {data['week_number']} ({data['date_range']})")
    title_run.bold = True
    title_run.font.size = Pt(16)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Total Portfolio Sales
    doc.add_heading("Total Portfolio Sales", level=2)
    p1 = doc.add_paragraph()
    p1.add_run(f"• KSH {data['total_sales']:,.0f}, marking a {data['total_change_percent']}% ")
    p1.add_run("decline" if data['total_change_percent'] < 0 else "increase").bold = True
    p1.add_run(f" from last week's KSH {data['last_week_sales']:,.0f}.")
    
    p2 = doc.add_paragraph()
    p2.add_run(f"• {data['total_sales_comment']}")
    
    # Performance by Sales Executives
    doc.add_heading("Performance by Sales Executives", level=2)
    
    for exec_data in data['executives']:
        doc.add_heading(exec_data['name'], level=3)
        p = doc.add_paragraph()
        p.add_run(f"• {exec_data['change_percent']}%: KSH {exec_data['current_sales']/1000:.2f}K (from KSH {exec_data['last_sales']/1000:.2f}K)\n")
        p.add_run(f"• {exec_data['comment']}")
    
    # Key Highlights
    doc.add_heading("Key Highlights", level=2)
    for highlight in data['highlights']:
        doc.add_paragraph(highlight, style="List Bullet")
    
    # Strategic Next Steps
    doc.add_heading("Strategic Next Steps", level=2)
    for i, step in enumerate(data['next_steps'], 1):
        doc.add_paragraph(f"{i}. {step}", style="List Number")
    
    return doc

def main():
    st.title("Weekly Sales Report Generator")
    
    # Input form
    with st.form("report_data"):
        col1, col2 = st.columns(2)
        with col1:
            week_number = st.number_input("Week Number", min_value=1, max_value=52, value=20)
            total_sales = st.number_input("Total Sales (KSH)", value=11808769)
            last_week_sales = st.number_input("Last Week Sales (KSH)", value=14583061)
            date_range = st.text_input("Date Range", value="May 9th - May 15th, 2025")
        
        with col2:
            total_change_percent = st.number_input("Total Change Percent", value=-19.0)
            total_sales_comment = st.text_area("Total Sales Comment", 
                                             value="A setback following a strong week—key for teams to stabilize and build consistency.")
        
        st.subheader("Sales Executives Performance")
        
        executives = [
            {"name": "Caroline", "current_sales": 1630000, "last_sales": 1690000, 
             "change_percent": -4, "comment": "Slight dip, but performance remains stable and strong."},
            {"name": "Edwin", "current_sales": 222000, "last_sales": 287000, 
             "change_percent": -22, "comment": "Coastal low season continues to impact sales — targeted support needed."},
            {"name": "Export", "current_sales": 300000, "last_sales": 2540000, 
             "change_percent": -88, "comment": "Significant drop post-Andes order — revisit client engagement plans for consistency."},
            {"name": "James", "current_sales": 631000, "last_sales": 978000, 
             "change_percent": -35, "comment": "Pipeline instability observed — close monitoring and follow-ups required."},
            {"name": "Janerose", "current_sales": 182000, "last_sales": 185000, 
             "change_percent": -2, "comment": "Stable performance amid seasonal trends — maintain proactive engagement."},
            {"name": "Josephine", "current_sales": 1100000, "last_sales": 847000, 
             "change_percent": 30, "comment": "Strong comeback — recent client management strategies proving effective."},
            {"name": "Mary", "current_sales": 3380000, "last_sales": 4150000, 
             "change_percent": -19, "comment": "Decline after peak orders — follow-ups and re-engagement critical for momentum."},
            {"name": "Moses", "current_sales": 4040000, "last_sales": 2870000, 
             "change_percent": 40, "comment": "Outstanding growth — strong upselling and client retention driving success."},
            {"name": "Nanyuki", "current_sales": 113000, "last_sales": 774000, 
             "change_percent": -85, "comment": "Expected biweekly drop — stay aligned to follow-up cycle for sustained performance."},
            {"name": "UPC (Upcountry)", "current_sales": 219000, "last_sales": 259000, 
             "change_percent": -15, "comment": "Moderate dip, though long-term trend remains upward — maintain momentum."},
        ]
        
        # Allow editing of executive data
        edited_executives = []
        for exec_data in executives:
            with st.expander(f"Edit {exec_data['name']}'s Data"):
                name = st.text_input("Name", value=exec_data['name'], key=f"name_{exec_data['name']}")
                current_sales = st.number_input("Current Sales", value=exec_data['current_sales'], key=f"current_{exec_data['name']}")
                last_sales = st.number_input("Last Week Sales", value=exec_data['last_sales'], key=f"last_{exec_data['name']}")
                change_percent = st.number_input("Change %", value=exec_data['change_percent'], key=f"change_{exec_data['name']}")
                comment = st.text_area("Comment", value=exec_data['comment'], key=f"comment_{exec_data['name']}")
                edited_executives.append({
                    "name": name,
                    "current_sales": current_sales,
                    "last_sales": last_sales,
                    "change_percent": change_percent,
                    "comment": comment
                })
        
        highlights = [
            f"Total sales declined by {abs(total_change_percent)}%, reversing previous week's growth.",
            "Moses and Josephine recorded strong gains — leading by example.",
            "Export, James, Mary, Edwin, and Nanyuki faced notable declines — each requiring specific strategic attention.",
            "Export's plunge highlights the risk of unsustained surges — prioritize continuity.",
            "Nanyuki's biweekly pattern continues — timing and cadence remain critical."
        ]
        
        edited_highlights = st.text_area("Key Highlights (one per line)", value="\n".join(highlights), height=150)
        
        next_steps = [
            "Sustain High Performers: Reinforce Moses and Josephine's winning strategies across the team.",
            "Stabilize Export Volatility: Evaluate client needs and engagement depth following large orders.",
            "Support Coastal Reps: Provide distribution and promotional support for Edwin and Janerose during the low season.",
            "Rebuild James & Mary's Pipeline: Focus on reactivation, follow-ups, and near-term conversions.",
            "Optimize Nanyuki's Cycle: Tighten biweekly rhythm to reduce sharp swings in performance."
        ]
        
        edited_next_steps = st.text_area("Strategic Next Steps (one per line)", value="\n".join(next_steps), height=150)
        
        submitted = st.form_submit_button("Generate Report")
    
    if submitted:
        # Prepare data for report generation
        report_data = {
            "week_number": week_number,
            "date_range": date_range,
            "total_sales": total_sales,
            "last_week_sales": last_week_sales,
            "total_change_percent": total_change_percent,
            "total_sales_comment": total_sales_comment,
            "executives": edited_executives,
            "highlights": edited_highlights.split("\n"),
            "next_steps": edited_next_steps.split("\n")
        }
        
        # Generate Word document
        doc = create_word_report(report_data)
        
        # Save to bytes buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        # Download button
        st.success("Report generated successfully!")
        st.download_button(
            label="Download Word Report",
            data=buffer,
            file_name=f"Weekly_Sales_Report_Week_{week_number}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

if __name__ == "__main__":
    main()